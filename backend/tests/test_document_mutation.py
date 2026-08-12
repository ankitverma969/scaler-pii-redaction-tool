from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from app.document import (
    DocumentMutationError,
    apply_replacements_to_block,
    apply_replacements_to_text,
    iter_text_blocks,
    load_docx,
)
from app.models import DetectedEntity, PIIType
from app.replacement import PlannedReplacement


def block_from_document(document: Document):
    return next(block for block in iter_text_blocks(document) if block.text)


def block_with_text(document: Document, expected: str):
    return next(block for block in iter_text_blocks(document) if block.text == expected)


def plan(
    text: str,
    value: str,
    replacement: str,
    pii_type: PIIType = PIIType.EMAIL,
    occurrence: int = 0,
) -> PlannedReplacement:
    cursor = 0
    start = -1
    for _ in range(occurrence + 1):
        start = text.index(value, cursor)
        cursor = start + len(value)
    entity = DetectedEntity(
        text=value,
        start=start,
        end=start + len(value),
        pii_type=pii_type,
        confidence=0.99,
        source="test",
    )
    return PlannedReplacement(entity=entity, replacement=replacement)


@pytest.mark.parametrize(
    "text,plans,expected",
    [
        ("abc", [("b", "X")], "aXc"),
        ("abc def", [("abc", "LONGER"), ("def", "x")], "LONGER x"),
        ("abcdef", [("ab", "X"), ("cd", "Y"), ("ef", "Z")], "XYZ"),
        ("abc", [("a", "X")], "Xbc"),
        ("abc", [("c", "X")], "abX"),
        ("abc", [("abc", "X")], "X"),
    ],
)
def test_apply_replacements_to_text_cases(text: str, plans, expected: str) -> None:
    planned = [plan(text, value, replacement) for value, replacement in plans]

    assert apply_replacements_to_text(text, planned) == expected


def test_apply_replacements_to_text_rejects_invalid_and_overlapping_plans() -> None:
    text = "abcdef"
    invalid = PlannedReplacement(
        entity=DetectedEntity("abc", 0, 3, PIIType.PERSON, 0.99, "test"),
        replacement="X",
    )
    object.__setattr__(invalid.entity, "end", 99)

    with pytest.raises(DocumentMutationError):
        apply_replacements_to_text(text, [invalid])

    overlapping = [
        PlannedReplacement(
            entity=DetectedEntity("abc", 0, 3, PIIType.PERSON, 0.99, "test"),
            replacement="X",
        ),
        PlannedReplacement(
            entity=DetectedEntity("cde", 2, 5, PIIType.COMPANY, 0.99, "test"),
            replacement="Y",
        ),
    ]
    with pytest.raises(DocumentMutationError):
        apply_replacements_to_text(text, overlapping)


def test_single_run_replacement_preserves_run_formatting() -> None:
    document = Document()
    run = document.add_paragraph().add_run("Contact Alice Smith today")
    run.bold = True
    run.italic = True

    block = block_from_document(document)
    applied = apply_replacements_to_block(
        block,
        [plan(block.text, "Alice Smith", "Meera Nair", PIIType.PERSON)],
    )

    assert applied == 1
    assert run.text == "Contact Meera Nair today"
    assert run.bold is True
    assert run.italic is True


def test_multi_run_replacement_preserves_prefix_suffix_and_run_objects() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Contact ")
    first = paragraph.add_run("Alice ")
    middle = paragraph.add_run("Smi")
    last = paragraph.add_run("th today")
    first.bold = True
    last.italic = True

    block = block_from_document(document)
    apply_replacements_to_block(
        block,
        [plan(block.text, "Alice Smith", "Meera Nair", PIIType.PERSON)],
    )

    assert [run.text for run in paragraph.runs] == [
        "Contact ",
        "Meera Nair",
        "",
        " today",
    ]
    assert first.bold is True
    assert middle.text == ""
    assert last.italic is True


def test_multiple_entities_in_same_run_apply_right_to_left() -> None:
    document = Document()
    paragraph = document.add_paragraph(
        "Email a@example.com or b@example.com today."
    )
    block = block_from_document(document)
    plans = [
        plan(block.text, "a@example.com", "first.long@example.com", PIIType.EMAIL),
        plan(block.text, "b@example.com", "b@e.co", PIIType.EMAIL),
    ]

    apply_replacements_to_block(block, plans)

    assert paragraph.runs[0].text == "Email first.long@example.com or b@e.co today."


def test_multiple_different_pii_types_in_one_run() -> None:
    document = Document()
    paragraph = document.add_paragraph(
        "Rahul Mehta | rahul@example.com | +91 98765 43210"
    )
    block = block_from_document(document)
    plans = [
        plan(block.text, "Rahul Mehta", "Meera Nair", PIIType.PERSON),
        plan(block.text, "rahul@example.com", "meera@example.com", PIIType.EMAIL),
        plan(block.text, "+91 98765 43210", "+91 90000 11111", PIIType.PHONE),
    ]

    apply_replacements_to_block(block, plans)

    assert paragraph.runs[0].text == (
        "Meera Nair | meera@example.com | +91 90000 11111"
    )


def test_entity_split_across_runs_replaced_once() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Email: ")
    paragraph.add_run("rah")
    paragraph.add_run("ul@")
    paragraph.add_run("example.")
    paragraph.add_run("com")

    block = block_from_document(document)
    apply_replacements_to_block(
        block,
        [plan(block.text, "rahul@example.com", "meera@example.com", PIIType.EMAIL)],
    )

    assert "".join(run.text for run in paragraph.runs) == "Email: meera@example.com"
    assert [run.text for run in paragraph.runs] == [
        "Email: ",
        "meera@example.com",
        "",
        "",
        "",
    ]


def test_replacements_apply_in_table_header_and_footer_blocks() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table Email: table@example.com"
    document.sections[0].header.paragraphs[0].text = "Header Phone: +91 98765 43210"
    document.sections[0].footer.paragraphs[0].text = "Footer DOB: 12/04/1990"

    table_block = block_with_text(document, "Table Email: table@example.com")
    header_block = block_with_text(document, "Header Phone: +91 98765 43210")
    footer_block = block_with_text(document, "Footer DOB: 12/04/1990")

    assert apply_replacements_to_block(
        table_block,
        [plan(table_block.text, "table@example.com", "safe@example.com", PIIType.EMAIL)],
    ) == 1
    assert apply_replacements_to_block(
        header_block,
        [plan(header_block.text, "+91 98765 43210", "+91 90000 11111", PIIType.PHONE)],
    ) == 1
    assert apply_replacements_to_block(
        footer_block,
        [plan(footer_block.text, "12/04/1990", "14/05/1988", PIIType.DOB)],
    ) == 1

    assert table.cell(0, 0).text == "Table Email: safe@example.com"
    assert document.sections[0].header.paragraphs[0].text == "Header Phone: +91 90000 11111"
    assert document.sections[0].footer.paragraphs[0].text == "Footer DOB: 14/05/1988"


def test_replacement_longer_shorter_and_adjacent_entities() -> None:
    document = Document()
    paragraph = document.add_paragraph("BobVery Long Example Company Limited")
    block = block_from_document(document)
    plans = [
        plan(block.text, "Bob", "Raghav Subramanian", PIIType.PERSON),
        plan(
            block.text,
            "Very Long Example Company Limited",
            "Acme Ltd",
            PIIType.COMPANY,
        ),
    ]

    apply_replacements_to_block(block, plans)

    assert paragraph.runs[0].text == "Raghav SubramanianAcme Ltd"


def test_span_boundaries_and_whole_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run("BC")
    paragraph.add_run("D")
    block = block_from_document(document)

    apply_replacements_to_block(block, [plan(block.text, "BC", "X", PIIType.PERSON)])

    assert [run.text for run in paragraph.runs] == ["A", "X", "D"]


def test_formatting_preservation_regression(tmp_path: Path) -> None:
    path = tmp_path / "formatting.docx"
    document = Document()
    paragraph = document.add_paragraph()
    normal_prefix = paragraph.add_run("Contact: ")
    first = paragraph.add_run("Rahul ")
    second = paragraph.add_run("Mehta")
    normal_suffix = paragraph.add_run(" for details.")
    first.bold = True
    first.font.name = "Arial"
    first.font.size = Pt(14)
    second.bold = True
    second.italic = True
    second.underline = True

    block = block_from_document(document)
    apply_replacements_to_block(
        block,
        [plan(block.text, "Rahul Mehta", "Meera Nair", PIIType.PERSON)],
    )
    document.save(path)
    reloaded = load_docx(path)

    assert [run.text for run in paragraph.runs] == [
        "Contact: ",
        "Meera Nair",
        "",
        " for details.",
    ]
    assert normal_prefix.text == "Contact: "
    assert normal_suffix.text == " for details."
    assert first.bold is True
    assert first.font.name == "Arial"
    assert first.font.size == Pt(14)
    assert second.bold is True
    assert second.italic is True
    assert second.underline is True
    assert "Meera Nair" in reloaded.paragraphs[0].text
