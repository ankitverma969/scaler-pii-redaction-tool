from collections import Counter

from docx import Document
from docx.enum.section import WD_SECTION

from app.detectors import StructuredPIIDetector
from app.detectors.address import AddressDetector
from app.detectors.company import CompanyDetector
from app.detectors.nlp import SpacyProvider
from app.detectors.person import PersonDetector
from app.detectors.semantic import SemanticPIIDetector
from app.document import SourceType, iter_text_blocks
from app.models import PIIType


def assert_entity_invariants(text, entities) -> None:
    previous = None
    for entity in entities:
        assert 0 <= entity.start < entity.end <= len(text)
        assert entity.text == text[entity.start : entity.end]
        assert 0.0 <= entity.confidence <= 1.0
        if previous is not None:
            assert (previous.start, previous.end, previous.pii_type.value, previous.source) <= (
                entity.start,
                entity.end,
                entity.pii_type.value,
                entity.source,
            )
        previous = entity


def types(entities):
    return [entity.pii_type for entity in entities]


def test_spacy_model_loads_and_contributes_person_candidate() -> None:
    provider = SpacyProvider()
    doc = provider.nlp("Barack Obama met the advisory team.")

    entities = PersonDetector(provider).detect(doc.text, doc)

    assert provider.nlp.meta["name"] == "core_web_sm"
    assert "Barack Obama" in [entity.text for entity in entities]


def test_spacy_provider_reuses_loaded_model() -> None:
    first = SpacyProvider().nlp
    second = SpacyProvider().nlp

    assert first is second


def test_person_context_positive_cases_and_exact_spans() -> None:
    text = (
        "Contact Person: Rahul Mehta\n"
        "Chief Financial Officer: Priya Nair\n"
        "Managing Director: Arjun Kulkarni\n"
        "Contact Person: Rahul Mehta / Priya Shah\n"
        "Company Secretary: Vikram S. Menon"
    )

    entities = PersonDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities if entity.source == "context_person"] == [
        "Rahul Mehta",
        "Priya Nair",
        "Arjun Kulkarni",
        "Rahul Mehta",
        "Priya Shah",
        "Vikram S. Menon",
    ]


def test_person_context_handles_uppercase_names_without_role_text() -> None:
    text = "Contact Person: RAHUL MEHTA\nBoard Committee: RISK MANAGEMENT COMMITTEE"

    entities = PersonDetector().detect(text)

    assert [entity.text for entity in entities] == ["RAHUL MEHTA"]


def test_person_role_titles_are_not_included() -> None:
    text = "Contact Person: Rahul Mehta, Company Secretary"

    entities = PersonDetector().detect(text)

    assert [entity.text for entity in entities if entity.pii_type == PIIType.PERSON] == [
        "Rahul Mehta"
    ]


def test_person_negative_domain_phrases() -> None:
    text = (
        "Board of Directors. General Information. Book Running Lead Managers. "
        "Risk Management Committee. Corporate Governance. Qualified Institutional Buyers."
    )

    assert PersonDetector().detect(text) == []


def test_person_negative_financial_and_governance_headings() -> None:
    text = (
        "Financial Statements. Independent Auditor Report. "
        "Corporate Social Responsibility Committee. Material Contracts."
    )

    assert PersonDetector().detect(text) == []


def test_company_legal_suffix_positive_cases() -> None:
    text = (
        "Aurora Industrial Systems Private Limited appointed Northstar Securities Limited. "
        "Example Finance LLP works with Summit Bank Limited."
    )

    entities = CompanyDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities if entity.source == "legal_suffix_company"] == [
        "Aurora Industrial Systems Private Limited",
        "Northstar Securities Limited",
        "Example Finance LLP",
        "Summit Bank Limited",
    ]


def test_company_detects_complete_span_with_punctuation() -> None:
    text = "(Example Wealth Management Limited), acted as advisor."

    entities = CompanyDetector().detect(text)

    assert [entity.text for entity in entities] == ["Example Wealth Management Limited"]


def test_company_negative_regulators_and_committees() -> None:
    text = (
        "Board of Directors. Audit Committee. Government of India. "
        "Registrar of Companies. Companies Act. SEBI ICDR Regulations. "
        "Risk Management Committee."
    )

    assert CompanyDetector().detect(text) == []


def test_company_rejects_act_and_regulation_phrases_with_company_words() -> None:
    text = (
        "Companies Act, 2013. Securities Contracts Regulation Act. "
        "The company law committee reviewed the draft."
    )

    assert CompanyDetector().detect(text) == []


def test_address_positive_indian_formats() -> None:
    text = (
        "Registered Office: 201, Tower 2, Example Business Centre, Baner, "
        "Pune - 411 045, Maharashtra, India; Telephone: +91 20 4505 3237"
    )

    entities = AddressDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "201, Tower 2, Example Business Centre, Baner, Pune - 411 045, Maharashtra, India"
    ]
    assert "Telephone" not in entities[0].text


def test_address_plot_and_multiline_formats() -> None:
    text = (
        "Plot No. 7, Industrial Area, Phase II, Village Example, Taluka Khed, "
        "Pune - 410501, Maharashtra, India"
    )
    multiline = (
        "801-804, Wing A, Building No. 3\n"
        "Inspire Business Complex\n"
        "Bandra East, Mumbai - 400 051\n"
        "Maharashtra, India"
    )

    assert [entity.text for entity in AddressDetector().detect(text)] == [text]
    assert [entity.text for entity in AddressDetector().detect(multiline)] == [multiline]


def test_address_negative_standalone_locations() -> None:
    text = (
        "The company operates in Pune. Sales increased in Maharashtra. "
        "The products are exported from India. Government of Maharashtra issued approval."
    )

    assert AddressDetector().detect(text) == []


def test_address_phone_email_mixed_block_keeps_semantic_separate() -> None:
    text = (
        "Registered Office: 201, Example Tower, Baner, Pune - 411045, Maharashtra, India; "
        "Telephone: +91 20 4505 3237; Email: office@example.com"
    )

    semantic_entities = SemanticPIIDetector().detect(text)
    structured_entities = StructuredPIIDetector().detect(text)

    assert Counter(types(semantic_entities)) == {PIIType.ADDRESS: 1}
    assert Counter(types(structured_entities)) == {PIIType.PHONE: 1, PIIType.EMAIL: 1}


def test_address_context_stops_before_email_phone_and_website() -> None:
    text = (
        "Corporate Office: 801-804, Wing A, Building No. 3, Inspire Business Complex, "
        "Bandra East, Mumbai - 400 051, Maharashtra, India; Email: office@example.com; "
        "Phone: +91 22 4444 1111; Website: www.example.com"
    )

    entities = AddressDetector().detect(text)

    assert [entity.text for entity in entities] == [
        "801-804, Wing A, Building No. 3, Inspire Business Complex, "
        "Bandra East, Mumbai - 400 051, Maharashtra, India"
    ]
    assert "office@example.com" not in entities[0].text
    assert "+91" not in entities[0].text
    assert "Website" not in entities[0].text


def test_semantic_detector_smoke_with_structured_detector() -> None:
    text = (
        "Contact Person: Meera Subramanian. "
        "Company: Aurora Industrial Systems Private Limited. "
        "Registered Office: 201, Example Tower, Baner, Pune - 411045, Maharashtra, India. "
        "Email: office@example.com; Phone: +91 98765 43210."
    )

    semantic = SemanticPIIDetector().detect(text)
    structured = StructuredPIIDetector().detect(text)

    assert {PIIType.PERSON, PIIType.COMPANY, PIIType.ADDRESS}.issubset(set(types(semantic)))
    assert {PIIType.EMAIL, PIIType.PHONE}.issubset(set(types(structured)))


def test_semantic_empty_text_is_safe() -> None:
    detector = SemanticPIIDetector()

    assert detector.detect("") == []
    assert detector.detect(" \n\t ") == []


def test_cross_run_person_maps_back_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Contact Person: Rahul ")
    paragraph.add_run("Meh")
    paragraph.add_run("ta")

    block = next(block for block in iter_text_blocks(document) if block.text)
    person = SemanticPIIDetector().detect(block.text)[0]
    run_slices = block.runs_for_span(person.start, person.end)

    assert person.text == "Rahul Mehta"
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_cross_run_company_maps_back_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Aurora Industrial ")
    paragraph.add_run("Systems Private ")
    paragraph.add_run("Limited")

    block = next(block for block in iter_text_blocks(document) if block.text)
    company = SemanticPIIDetector().detect(block.text)[0]
    run_slices = block.runs_for_span(company.start, company.end)

    assert company.text == "Aurora Industrial Systems Private Limited"
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_cross_run_address_maps_back_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Registered Office: 201, Example Tower, ")
    paragraph.add_run("Baner, Pune - ")
    paragraph.add_run("411045, Maharashtra, India")

    block = next(block for block in iter_text_blocks(document) if block.text)
    address = SemanticPIIDetector().detect(block.text)[0]
    run_slices = block.runs_for_span(address.start, address.end)

    assert address.pii_type == PIIType.ADDRESS
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_semantic_table_detection() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "Contact Person: Ananya Reddy"
    table.cell(1, 0).text = "Company: Northstar Securities Limited"
    table.cell(2, 0).text = "Address: 201, Example Tower, Baner, Pune - 411045, Maharashtra, India"

    detector = SemanticPIIDetector()
    table_blocks = [
        block for block in iter_text_blocks(document) if block.source_type == SourceType.TABLE
    ]

    assert PIIType.PERSON in types(detector.detect(table_blocks[0].text))
    assert PIIType.COMPANY in types(detector.detect(table_blocks[1].text))
    assert PIIType.ADDRESS in types(detector.detect(table_blocks[2].text))


def test_semantic_shared_header_detection_once() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Company: Summit Bank Limited"
    document.add_section(WD_SECTION.NEW_PAGE)

    detector = SemanticPIIDetector()
    entities = [
        entity
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.HEADER
        for entity in detector.detect(block.text)
    ]

    assert [entity.text for entity in entities] == ["Summit Bank Limited"]
