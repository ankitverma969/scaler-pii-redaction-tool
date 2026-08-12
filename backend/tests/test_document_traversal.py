from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches

from app.document import (
    DocxNotFoundError,
    InvalidDocxError,
    SourceType,
    UnsupportedDocumentTypeError,
    collect_document_statistics,
    iter_text_blocks,
    load_docx,
)
from app.document.text_map import run_span_offsets


def save_docx(document: Document, path: Path) -> Path:
    document.save(path)
    return path


def non_empty_blocks(document: Document):
    return [block for block in iter_text_blocks(document) if block.text]


def test_simple_paragraph_logical_text() -> None:
    document = Document()
    document.add_paragraph("Hello world")

    blocks = non_empty_blocks(document)

    assert blocks[0].text == "Hello world"
    assert blocks[0].source_type == SourceType.BODY


def test_multi_run_paragraph_offsets() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("John ")
    paragraph.add_run("Do")
    paragraph.add_run("e")

    block = non_empty_blocks(document)[0]

    assert block.text == "John Doe"
    assert run_span_offsets(block.run_spans) == [
        (0, 0, 5),
        (1, 5, 7),
        (2, 7, 8),
    ]


def test_span_crossing_runs_maps_to_all_affected_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("John ")
    paragraph.add_run("Do")
    paragraph.add_run("e")

    block = non_empty_blocks(document)[0]
    slices = block.runs_for_span(0, 8)

    assert [(s.run_index, s.start, s.end, s.start_in_run, s.end_in_run) for s in slices] == [
        (0, 0, 5, 0, 5),
        (1, 5, 7, 0, 2),
        (2, 7, 8, 0, 1),
    ]


def test_partial_span_inside_single_run() -> None:
    document = Document()
    document.add_paragraph("Hello world")

    block = non_empty_blocks(document)[0]
    slices = block.runs_for_span(6, 11)

    assert [(s.run_index, s.start, s.end, s.start_in_run, s.end_in_run) for s in slices] == [
        (0, 6, 11, 6, 11)
    ]


def test_empty_runs_do_not_corrupt_offsets() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run("")
    paragraph.add_run("B")

    block = non_empty_blocks(document)[0]

    assert block.text == "AB"
    assert run_span_offsets(block.run_spans) == [
        (0, 0, 1),
        (1, 1, 1),
        (2, 1, 2),
    ]
    assert [(s.run_index, s.start, s.end) for s in block.runs_for_span(0, 2)] == [
        (0, 0, 1),
        (2, 1, 2),
    ]


def test_tabs_and_line_breaks_are_preserved_in_logical_text() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("A")
    run.add_tab()
    run.add_text("B")
    run.add_break()
    run.add_text("C")

    block = non_empty_blocks(document)[0]

    assert block.text == "A\tB\nC"


def test_invalid_span_raises_value_error() -> None:
    document = Document()
    document.add_paragraph("Hello")

    block = non_empty_blocks(document)[0]

    with pytest.raises(ValueError):
        block.runs_for_span(0, 99)


def test_table_traversal_emits_cell_paragraphs_once() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"

    table_blocks = [
        block for block in iter_text_blocks(document) if block.source_type == SourceType.TABLE
    ]

    assert [block.text for block in table_blocks if block.text] == ["A1", "A2", "B1", "B2"]
    assert len({block.block_id for block in table_blocks}) == len(table_blocks)


def test_nested_table_traversal() -> None:
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    outer.cell(0, 0).text = "Outer"
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested"

    texts = [block.text for block in iter_text_blocks(document) if block.text]

    assert "Outer" in texts
    assert "Nested" in texts


def test_merged_table_cell_is_not_emitted_twice() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged"

    table_texts = [
        block.text
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.TABLE and block.text
    ]

    assert table_texts == ["Merged"]


def test_header_text_is_emitted() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header text"

    header_texts = [
        block.text
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.HEADER and block.text
    ]

    assert header_texts == ["Header text"]


def test_shared_header_part_is_emitted_once() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Shared header"
    document.add_section(WD_SECTION.NEW_PAGE)

    header_texts = [
        block.text
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.HEADER and block.text == "Shared header"
    ]

    assert header_texts == ["Shared header"]


def test_footer_text_is_emitted() -> None:
    document = Document()
    document.sections[0].footer.paragraphs[0].text = "Footer text"

    footer_texts = [
        block.text
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.FOOTER and block.text
    ]

    assert footer_texts == ["Footer text"]


def test_header_table_is_traversed() -> None:
    document = Document()
    table = document.sections[0].header.add_table(rows=1, cols=1, width=Inches(2))
    table.cell(0, 0).text = "Header table"

    header_texts = [
        block.text
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.HEADER and block.text
    ]

    assert "Header table" in header_texts


def test_deterministic_block_ids_after_reload(tmp_path: Path) -> None:
    path = tmp_path / "stable.docx"
    document = Document()
    document.add_paragraph("Stable")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    save_docx(document, path)

    first = [block.block_id for block in iter_text_blocks(load_docx(path))]
    second = [block.block_id for block in iter_text_blocks(load_docx(path))]

    assert first == second


def test_loader_rejects_missing_file(tmp_path: Path) -> None:
    with pytest.raises(DocxNotFoundError):
        load_docx(tmp_path / "missing.docx")


def test_loader_rejects_non_docx_extension(tmp_path: Path) -> None:
    path = tmp_path / "not-docx.txt"
    path.write_text("not a docx")

    with pytest.raises(UnsupportedDocumentTypeError):
        load_docx(path)


def test_loader_rejects_invalid_docx(tmp_path: Path) -> None:
    path = tmp_path / "invalid.docx"
    path.write_text("not a real docx")

    with pytest.raises(InvalidDocxError):
        load_docx(path)


def test_statistics_are_safe_aggregate_counts() -> None:
    document = Document()
    document.add_paragraph("Body")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    document.sections[0].footer.paragraphs[0].text = "Footer"

    stats = collect_document_statistics(document)

    assert stats.section_count == 1
    assert stats.table_count == 1
    assert stats.text_block_count >= 3
    assert stats.footer_part_count >= 1
