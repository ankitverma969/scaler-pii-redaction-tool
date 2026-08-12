from collections import Counter

from docx import Document
from docx.enum.section import WD_SECTION

from app.detectors import (
    CreditCardDetector,
    DOBDetector,
    EmailDetector,
    IPAddressDetector,
    PhoneDetector,
    SSNDetector,
    StructuredPIIDetector,
)
from app.detectors.validators import (
    is_valid_date,
    is_valid_ipv4,
    is_valid_ssn,
    normalize_digits,
    passes_luhn,
)
from app.document import SourceType, iter_text_blocks
from app.models import PIIType


def assert_entity_invariants(text, entities) -> None:
    previous = None
    for entity in entities:
        assert 0 <= entity.start < entity.end <= len(text)
        assert entity.text == text[entity.start : entity.end]
        assert 0.0 <= entity.confidence <= 1.0
        if previous is not None:
            assert (previous.start, previous.end, previous.pii_type.value, previous.source) <= (
                entity.start,
                entity.end,
                entity.pii_type.value,
                entity.source,
            )
        previous = entity


def types(entities):
    return [entity.pii_type for entity in entities]


def test_email_detector_positive_cases_and_exact_spans() -> None:
    text = (
        "Email: john.doe@example.com, backup foo+bar@mail.example.co.in; "
        "(finance-team@example.org)"
    )

    entities = EmailDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "john.doe@example.com",
        "foo+bar@mail.example.co.in",
        "finance-team@example.org",
    ]
    assert all(entity.pii_type == PIIType.EMAIL for entity in entities)


def test_email_detector_rejects_malformed_addresses() -> None:
    text = "Bad: john@ @example.com john@example john..doe@example.com"

    assert EmailDetector().detect(text) == []


def test_phone_detector_positive_cases() -> None:
    text = (
        "Mobile +91 9876543210; backup 9876543210; spaced +91 98765 43210; "
        "hyphen +91-9876543210; Telephone: +91 20 4505 3237; Tel: +91-20-26234000"
    )

    entities = PhoneDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "+91 9876543210",
        "9876543210",
        "+91 98765 43210",
        "+91-9876543210",
        "+91 20 4505 3237",
        "+91-20-26234000",
    ]
    assert all(entity.pii_type == PIIType.PHONE for entity in entities)


def test_phone_detector_requires_context_for_landline_like_numbers() -> None:
    assert PhoneDetector().detect("Reference 20 4505 3237 is not labeled.") == []
    entities = PhoneDetector().detect("Telephone: 20 4505 3237")
    assert [entity.text for entity in entities] == ["20 4505 3237"]


def test_phone_detector_hard_negatives() -> None:
    text = (
        "CIN: U28129PN1979PLC141032 DIN: 00135070 "
        "SEBI Registration: INM000013004 PIN 411045 "
        "Amount INR 7100.00 million and 26,437,554 shares. "
        "Years 20252024 and date March 31, 2025."
    )

    assert PhoneDetector().detect(text) == []


def test_ssn_detector_validation() -> None:
    text = "Valid 123-45-6789 invalid 000-45-6789 666-45-6789 900-45-6789 123-00-6789 123-45-0000 123456789"

    entities = SSNDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == ["123-45-6789"]
    assert is_valid_ssn("123-45-6789")
    assert not is_valid_ssn("000-45-6789")
    assert not is_valid_ssn("666-45-6789")
    assert not is_valid_ssn("900-45-6789")
    assert not is_valid_ssn("123-00-6789")
    assert not is_valid_ssn("123-45-0000")


def test_ssn_boundary_protection() -> None:
    assert SSNDetector().detect("ABC123-45-6789XYZ") == []


def test_credit_card_detector_and_luhn_validator() -> None:
    text = (
        "Cards: 4111111111111111, 4111 1111 1111 1111, "
        "4111-1111-1111-1111. Invalid 4111 1111 1111 1112."
    )

    entities = CreditCardDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "4111111111111111",
        "4111 1111 1111 1111",
        "4111-1111-1111-1111",
    ]
    assert passes_luhn("4111 1111 1111 1111")
    assert not passes_luhn("4111 1111 1111 1112")
    assert normalize_digits("4111-1111") == "41111111"


def test_credit_card_detector_hard_negatives() -> None:
    text = (
        "Shares 26,437,554 and phone +91 98765 43210 and amount 7100.00. "
        "Registration U28129PN1979PLC141032 and random 1234567890123456."
    )

    assert CreditCardDetector().detect(text) == []


def test_ip_detector_ipv4_positive_and_negative_cases() -> None:
    text = "IPs: 192.0.2.10, 198.51.100.20 and 203.0.113.15. Bad 999.1.1.1 256.0.0.1 1.2.3 v1.2.3.4x"

    entities = IPAddressDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "192.0.2.10",
        "198.51.100.20",
        "203.0.113.15",
    ]
    assert is_valid_ipv4("192.0.2.10")
    assert not is_valid_ipv4("256.0.0.1")
    assert IPAddressDetector().detect("U192.0.2.1TEST") == []


def test_dob_detector_contextual_positive_cases() -> None:
    text = (
        "DOB: 12/04/1990. Date of Birth: 12-04-1990. "
        "Born on 12 April 1990. DOB: 1990-04-12. "
        "Date of Birth: April 12, 1990."
    )

    entities = DOBDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert [entity.text for entity in entities] == [
        "12/04/1990",
        "12-04-1990",
        "12 April 1990",
        "1990-04-12",
        "April 12, 1990",
    ]
    assert all(entity.pii_type == PIIType.DOB for entity in entities)


def test_dob_detector_rejects_ordinary_dates_and_invalid_dates() -> None:
    text = (
        "Agreement dated December 9, 2025. Incorporated on July 30, 1979. "
        "Fiscal year ended March 31, 2025. Offer opens December 16, 2025. "
        "DOB: 31/02/2020. Date of Birth: 99/99/1990. DOB: 2025-25-90."
    )

    assert DOBDetector().detect(text) == []
    assert not is_valid_date("31/02/2020")
    assert not is_valid_date("99/99/1990")
    assert not is_valid_date("2025-25-90")


def test_structured_detector_multiple_entities_and_no_person_detection() -> None:
    text = (
        "Contact: Jane Example\n"
        "Email: jane@example.com\n"
        "Telephone: +91 20 4505 3237\n"
        "Backup primary@example.com or backup@example.org, phone +91 98765 43210."
    )

    entities = StructuredPIIDetector().detect(text)

    assert_entity_invariants(text, entities)
    assert Counter(types(entities)) == {
        PIIType.EMAIL: 3,
        PIIType.PHONE: 2,
    }
    assert PIIType.PERSON not in types(entities)


def test_structured_detector_empty_and_whitespace_text() -> None:
    detector = StructuredPIIDetector()

    assert detector.detect("") == []
    assert detector.detect(" \n\t ") == []


def test_structured_detector_hard_negative_block() -> None:
    text = (
        "CIN: U28129PN1979PLC141032\n"
        "DIN: 00135070\n"
        "SEBI Registration: INM000013004\n"
        "INR 7,100.00 million\n"
        "26,437,554 shares\n"
        "March 31, 2025\n"
        "PIN 411045\n"
        "Section 32\n"
        "Regulation 6(1)\n"
    )

    entities = StructuredPIIDetector().detect(text)

    assert [entity.pii_type for entity in entities] == []


def test_cross_run_email_detection_maps_back_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("john.")
    paragraph.add_run("doe@")
    paragraph.add_run("example.com")

    block = next(block for block in iter_text_blocks(document) if block.text)
    entities = StructuredPIIDetector().detect(block.text)
    email = entities[0]
    run_slices = block.runs_for_span(email.start, email.end)

    assert email.text == "john.doe@example.com"
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_cross_run_phone_detection_maps_back_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Telephone: +91 ")
    paragraph.add_run("98765")
    paragraph.add_run(" 43210")

    block = next(block for block in iter_text_blocks(document) if block.text)
    entities = StructuredPIIDetector().detect(block.text)
    phone = entities[0]
    run_slices = block.runs_for_span(phone.start, phone.end)

    assert phone.text == "+91 98765 43210"
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_table_textblock_detection() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Email: table@example.com"
    table.cell(1, 0).text = "Phone: +91 98765 43210"

    detector = StructuredPIIDetector()
    entities = [
        entity
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.TABLE
        for entity in detector.detect(block.text)
    ]

    assert Counter(types(entities)) == {PIIType.EMAIL: 1, PIIType.PHONE: 1}


def test_shared_header_detection_once() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Email: header@example.com"
    document.add_section(WD_SECTION.NEW_PAGE)

    detector = StructuredPIIDetector()
    entities = [
        entity
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.HEADER
        for entity in detector.detect(block.text)
    ]

    assert [entity.text for entity in entities] == ["header@example.com"]
