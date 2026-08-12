from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches

from app.document import collect_document_statistics, iter_text_blocks, load_docx
from app.models import PIIType
from app.redaction import RedactionEngine


def all_text(path: Path) -> str:
    document = load_docx(path)
    return "\n".join(block.text for block in iter_text_blocks(document))


def block_texts(path: Path) -> dict[str, str]:
    document = load_docx(path)
    return {block.block_id: block.text for block in iter_text_blocks(document)}


def create_all_nine_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph(
        "Contact Person: Rahul Mehta\n"
        "Company: Aurora Industrial Systems Private Limited\n"
        "Registered Office: 10 Example Road, Baner, Pune - 411045, Maharashtra, India\n"
        "Email: rahul@example.com\n"
        "Telephone: +91 98765 43210\n"
        "DOB: 12/04/1990\n"
        "Server IP: 192.0.2.15\n"
        "SSN: 123-45-6789\n"
        "Test card: 4111 1111 1111 1111"
    )
    document.save(path)
    return path


def test_engine_redacts_all_nine_categories_and_validates_output(tmp_path: Path) -> None:
    source = create_all_nine_docx(tmp_path / "all-nine.docx")
    output = tmp_path / "all-nine-redacted.docx"

    result = RedactionEngine().redact(source, output, seed=42)

    assert result.output_path == output
    assert result.validation.success
    assert result.total_entities == result.total_planned == result.total_applied
    assert set(result.counts_by_type) == {pii_type.value for pii_type in PIIType}
    assert all(result.counts_by_type[pii_type.value] >= 1 for pii_type in PIIType)
    output_text = all_text(output)
    for original in [
        "Rahul Mehta",
        "rahul@example.com",
        "+91 98765 43210",
        "123-45-6789",
        "4111 1111 1111 1111",
        "192.0.2.15",
    ]:
        assert original not in output_text


def test_engine_preserves_tables_headers_footers_and_sections(tmp_path: Path) -> None:
    source = tmp_path / "structured.docx"
    output = tmp_path / "structured-redacted.docx"
    document = Document()
    document.add_paragraph("Email: body@example.com")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Company: Northstar Securities Limited"
    nested = table.cell(1, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Email: nested@example.com"
    document.sections[0].header.paragraphs[0].text = "Email: header@example.com"
    document.sections[0].footer.paragraphs[0].text = "DOB: 12/04/1990"
    header_table = document.sections[0].header.add_table(
        rows=1, cols=1, width=Inches(2)
    )
    header_table.cell(0, 0).text = "Email: table-header@example.com"
    document.add_section(WD_SECTION.NEW_PAGE)
    document.save(source)
    source_stats = collect_document_statistics(load_docx(source))

    result = RedactionEngine().redact(source, output, seed=42)
    output_stats = collect_document_statistics(load_docx(output))

    assert result.validation.success
    assert output_stats.section_count == source_stats.section_count
    assert output_stats.table_count == source_stats.table_count
    assert output_stats.nested_table_count == source_stats.nested_table_count
    assert output_stats.header_part_count == source_stats.header_part_count
    assert output_stats.footer_part_count == source_stats.footer_part_count
    output_text = all_text(output)
    assert "body@example.com" not in output_text
    assert "nested@example.com" not in output_text
    assert "header@example.com" not in output_text
    assert "table-header@example.com" not in output_text


def test_engine_same_seed_produces_same_logical_output(tmp_path: Path) -> None:
    source = create_all_nine_docx(tmp_path / "deterministic.docx")
    output_a = tmp_path / "a.docx"
    output_b = tmp_path / "b.docx"

    result_a = RedactionEngine().redact(source, output_a, seed=42)
    result_b = RedactionEngine().redact(source, output_b, seed=42)

    assert result_a.counts_by_type == result_b.counts_by_type
    assert block_texts(output_a) == block_texts(output_b)


def test_engine_different_seed_changes_synthetic_output(tmp_path: Path) -> None:
    source = create_all_nine_docx(tmp_path / "different-seed.docx")
    output_a = tmp_path / "seed-42.docx"
    output_b = tmp_path / "seed-43.docx"

    RedactionEngine().redact(source, output_a, seed=42)
    RedactionEngine().redact(source, output_b, seed=43)

    assert block_texts(output_a) != block_texts(output_b)
