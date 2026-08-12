from __future__ import annotations

import io
import threading
import time
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from fastapi.testclient import TestClient

from app.document import iter_text_blocks, load_docx
from app.jobs import JobManager
from app.jobs.models import empty_counts
from app.main import create_app


def docx_bytes(text: str) -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(stream)
    return stream.getvalue()


def write_docx(path: Path, text: str) -> Path:
    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return path


def client_with_manager(manager: JobManager) -> TestClient:
    return TestClient(create_app(job_manager=manager))


def poll_completed(client: TestClient, job_id: str, timeout: float = 30.0) -> dict:
    deadline = time.time() + timeout
    last_status = None
    while time.time() < deadline:
        response = client.get(f"/api/redactions/{job_id}")
        assert response.status_code == 200
        payload = response.json()
        last_status = payload["status"]
        if last_status in {"COMPLETED", "FAILED"}:
            return payload
        time.sleep(0.1)
    raise AssertionError(f"job did not finish; last status={last_status}")


def upload(client: TestClient, content: bytes, filename: str = "sample.docx", seed: int = 42):
    return client.post(
        "/api/redactions",
        data={"seed": str(seed)},
        files={
            "file": (
                filename,
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def test_valid_upload_status_download_and_delete(tmp_path: Path) -> None:
    manager = JobManager(temp_root=tmp_path / "jobs")
    secret = "private.person@example.com"

    with client_with_manager(manager) as client:
        accepted = upload(client, docx_bytes(f"Email: {secret}"))

        assert accepted.status_code == 202
        accepted_payload = accepted.json()
        assert accepted_payload["job_id"]
        assert accepted_payload["status"] in {"QUEUED", "PROCESSING", "COMPLETED"}
        assert secret not in accepted.text

        completed = poll_completed(client, accepted_payload["job_id"])

        assert completed["status"] == "COMPLETED"
        assert completed["total_entities"] == 1
        assert set(completed["counts"]) == set(empty_counts())
        assert completed["counts"]["EMAIL"] == 1
        assert completed["total_entities"] == sum(completed["counts"].values())
        assert completed["download_available"] is True
        assert secret not in str(completed)

        downloaded = client.get(f"/api/redactions/{accepted_payload['job_id']}/download")

        assert downloaded.status_code == 200
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in downloaded.headers["content-type"]
        assert zipfile.is_zipfile(io.BytesIO(downloaded.content))
        downloaded_path = tmp_path / "downloaded.docx"
        downloaded_path.write_bytes(downloaded.content)
        output_text = "\n".join(
            block.text for block in iter_text_blocks(load_docx(downloaded_path))
        )
        assert secret not in output_text

        internal_job = manager._internal_job(accepted_payload["job_id"])
        assert not internal_job.input_path.exists()
        assert internal_job.output_path.exists()
        assert client.delete(f"/api/redactions/{accepted_payload['job_id']}").json() == {
            "job_id": accepted_payload["job_id"],
            "deleted": True,
        }
        assert not internal_job.temp_dir.exists()
        assert client.get(f"/api/redactions/{accepted_payload['job_id']}").status_code == 404


def test_invalid_extension_fake_docx_empty_and_oversized_uploads(tmp_path: Path) -> None:
    manager = JobManager(temp_root=tmp_path / "jobs", max_upload_size_mb=1)
    with client_with_manager(manager) as client:
        invalid_ext = upload(client, b"notes", filename="notes.txt")
        fake_docx = upload(client, b"not a docx", filename="fake.docx")
        empty = upload(client, b"", filename="empty.docx")
        oversized = upload(client, b"x" * (1024 * 1024 + 1), filename="large.docx")

    assert invalid_ext.status_code == 400
    assert invalid_ext.json()["detail"]["code"] == "UNSUPPORTED_FILE_TYPE"
    assert fake_docx.status_code == 400
    assert fake_docx.json()["detail"]["code"] == "INVALID_DOCX"
    assert empty.status_code == 400
    assert empty.json()["detail"]["code"] == "EMPTY_FILE"
    assert oversized.status_code == 413
    assert oversized.json()["detail"]["code"] == "UPLOAD_TOO_LARGE"


def test_unknown_job_endpoints_return_404(tmp_path: Path) -> None:
    manager = JobManager(temp_root=tmp_path / "jobs")
    with client_with_manager(manager) as client:
        assert client.get("/api/redactions/missing").status_code == 404
        assert client.get("/api/redactions/missing/download").status_code == 404
        assert client.delete("/api/redactions/missing").status_code == 404


def test_download_before_completion_and_active_delete_return_409(tmp_path: Path) -> None:
    started = threading.Event()
    release = threading.Event()

    class SlowEngine:
        def redact(self, input_path, output_path, seed=42):
            started.set()
            release.wait(timeout=5)
            Path(output_path).write_bytes(Path(input_path).read_bytes())
            counts = empty_counts()
            counts["EMAIL"] = 1
            return SimpleNamespace(
                total_entities=1,
                counts_by_type=counts,
                duration_seconds=0.1,
            )

    manager = JobManager(temp_root=tmp_path / "jobs", engine_factory=SlowEngine)
    with client_with_manager(manager) as client:
        accepted = upload(client, docx_bytes("Email: queued@example.com"))
        job_id = accepted.json()["job_id"]
        assert started.wait(timeout=5)

        assert client.get(f"/api/redactions/{job_id}/download").status_code == 409
        assert client.delete(f"/api/redactions/{job_id}").status_code == 409

        release.set()
        assert poll_completed(client, job_id)["status"] == "COMPLETED"


def test_engine_failure_is_safe_and_not_downloadable(tmp_path: Path) -> None:
    class FailingEngine:
        def redact(self, input_path, output_path, seed=42):
            raise RuntimeError("private.person@example.com exploded")

    manager = JobManager(temp_root=tmp_path / "jobs", engine_factory=FailingEngine)
    with client_with_manager(manager) as client:
        accepted = upload(client, docx_bytes("Email: private.person@example.com"))
        job_id = accepted.json()["job_id"]
        failed = poll_completed(client, job_id)

        assert failed["status"] == "FAILED"
        assert failed["error"] == {
            "code": "REDACTION_FAILED",
            "message": "The document could not be processed.",
        }
        assert "private.person@example.com" not in str(failed)
        assert str(manager.temp_root) not in str(failed)
        assert client.get(f"/api/redactions/{job_id}/download").status_code == 409
        internal_job = manager._internal_job(job_id)
        assert not internal_job.input_path.exists()
        assert not internal_job.output_path.exists()


def test_ttl_cleanup_removes_expired_completed_jobs(tmp_path: Path) -> None:
    current = datetime(2026, 1, 1, tzinfo=timezone.utc)

    def now() -> datetime:
        return current

    manager = JobManager(temp_root=tmp_path / "jobs", ttl_minutes=1, now=now)
    with client_with_manager(manager) as client:
        accepted = upload(client, docx_bytes("Email: ttl@example.com"))
        job_id = accepted.json()["job_id"]
        completed = poll_completed(client, job_id)
        assert completed["status"] == "COMPLETED"
        internal_job = manager._internal_job(job_id)

        current = current + timedelta(minutes=2)
        assert client.get(f"/api/redactions/{job_id}").status_code == 404
        assert not internal_job.temp_dir.exists()


def test_concurrency_limit_with_slow_engine(tmp_path: Path) -> None:
    active = 0
    max_active = 0
    lock = threading.Lock()
    release = threading.Event()
    first_started = threading.Event()

    class SlowEngine:
        def redact(self, input_path, output_path, seed=42):
            nonlocal active, max_active
            with lock:
                active += 1
                max_active = max(max_active, active)
            first_started.set()
            release.wait(timeout=5)
            Path(output_path).write_bytes(Path(input_path).read_bytes())
            with lock:
                active -= 1
            return SimpleNamespace(
                total_entities=0,
                counts_by_type=empty_counts(),
                duration_seconds=0.1,
            )

    manager = JobManager(
        temp_root=tmp_path / "jobs", max_workers=1, engine_factory=SlowEngine
    )
    with client_with_manager(manager) as client:
        first = upload(client, docx_bytes("Email: first@example.com")).json()["job_id"]
        assert first_started.wait(timeout=5)
        second = upload(client, docx_bytes("Email: second@example.com")).json()["job_id"]
        release.set()
        assert poll_completed(client, first)["status"] == "COMPLETED"
        assert poll_completed(client, second)["status"] == "COMPLETED"

    assert max_active == 1


def test_filename_sanitization_and_job_isolation(tmp_path: Path) -> None:
    manager = JobManager(temp_root=tmp_path / "jobs")
    with client_with_manager(manager) as client:
        first = upload(
            client,
            docx_bytes("Email: first-private@example.com"),
            filename="../../private.docx",
            seed=42,
        ).json()["job_id"]
        second = upload(
            client,
            docx_bytes("Email: second-private@example.com"),
            filename="..\\..\\other.docx",
            seed=43,
        ).json()["job_id"]
        assert first != second
        assert poll_completed(client, first)["status"] == "COMPLETED"
        assert poll_completed(client, second)["status"] == "COMPLETED"
        first_job = manager._internal_job(first)
        second_job = manager._internal_job(second)

        assert first_job.temp_dir.parent == manager.temp_root
        assert second_job.temp_dir.parent == manager.temp_root
        assert first_job.temp_dir != second_job.temp_dir
        assert first_job.download_filename == "private_Redacted.docx"
        assert second_job.download_filename == "other_Redacted.docx"
        first_download = client.get(f"/api/redactions/{first}/download").content
        second_download = client.get(f"/api/redactions/{second}/download").content
        assert first_download != second_download
        assert client.delete(f"/api/redactions/{first}").status_code == 200
        assert client.get(f"/api/redactions/{first}").status_code == 404
        assert client.get(f"/api/redactions/{second}").status_code == 200
        assert client.get(f"/api/redactions/{second}/download").status_code == 200


def test_seed_determinism_for_api_jobs(tmp_path: Path) -> None:
    manager = JobManager(temp_root=tmp_path / "jobs")
    content = docx_bytes("Email: seed@example.com")
    with client_with_manager(manager) as client:
        job_a = upload(client, content, seed=42).json()["job_id"]
        job_b = upload(client, content, seed=42).json()["job_id"]
        job_c = upload(client, content, seed=43).json()["job_id"]
        assert poll_completed(client, job_a)["status"] == "COMPLETED"
        assert poll_completed(client, job_b)["status"] == "COMPLETED"
        assert poll_completed(client, job_c)["status"] == "COMPLETED"

        def logical_text(job_id: str) -> str:
            response = client.get(f"/api/redactions/{job_id}/download")
            path = tmp_path / f"{job_id}.docx"
            path.write_bytes(response.content)
            return "\n".join(block.text for block in iter_text_blocks(load_docx(path)))

        assert logical_text(job_a) == logical_text(job_b)
        assert logical_text(job_a) != logical_text(job_c)


def test_stale_temp_root_cleanup_on_startup(tmp_path: Path) -> None:
    stale = tmp_path / "jobs" / "stale-job"
    stale.mkdir(parents=True)
    (stale / "input.docx").write_bytes(b"private")
    manager = JobManager(temp_root=tmp_path / "jobs")

    with client_with_manager(manager) as client:
        assert client.get("/api/health").status_code == 200

    assert not stale.exists()
