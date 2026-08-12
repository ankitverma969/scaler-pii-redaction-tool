from collections import Counter

from docx import Document
from docx.enum.section import WD_SECTION

from app.detectors.unified import PIIDetector
from app.document import SourceType, iter_text_blocks
from app.models import PIIType


def types(entities):
    return [entity.pii_type for entity in entities]


def assert_final_invariants(text: str, entities) -> None:
    previous = None
    for entity in entities:
        assert 0 <= entity.start < entity.end <= len(text)
        assert entity.text == text[entity.start : entity.end]
        assert 0.0 <= entity.confidence <= 1.0
        if previous is not None:
            assert previous.end <= entity.start
        previous = entity


def test_unified_detector_returns_all_nine_categories() -> None:
    text = (
        "Contact Person: Rahul Mehta\n"
        "Company: Aurora Industrial Systems Private Limited\n"
        "Registered Office: 10 Example Road, Baner, Pune - 411045, Maharashtra, India\n"
        "Email: rahul@example.com\n"
        "Telephone: +91 98765 43210\n"
        "DOB: 12/04/1990\n"
        "Server IP: 192.0.2.15\n"
        "SSN: 123-45-6789\n"
        "Test card: 4111 1111 1111 1111"
    )

    entities = PIIDetector().detect(text)

    assert_final_invariants(text, entities)
    assert set(types(entities)) == set(PIIType)


def test_unified_detector_combined_synthetic_example_no_overlap() -> None:
    text = (
        "Contact Person: Rahul Mehta\n"
        "Company: Aurora Industrial Systems Private Limited\n"
        "Registered Office: 10 Example Road, Baner, Pune - 411045, Maharashtra, India\n"
        "Email: rahul@example.com\n"
        "Telephone: +91 98765 43210\n"
        "DOB: 12/04/1990\n"
        "Server IP: 192.0.2.15"
    )

    entities = PIIDetector().detect(text)

    assert_final_invariants(text, entities)
    assert {
        PIIType.PERSON,
        PIIType.COMPANY,
        PIIType.ADDRESS,
        PIIType.EMAIL,
        PIIType.PHONE,
        PIIType.DOB,
        PIIType.IP_ADDRESS,
    }.issubset(set(types(entities)))


def test_hard_negative_unified_block() -> None:
    text = (
        "GENERAL INFORMATION\n"
        "RISK FACTORS\n"
        "BOARD OF DIRECTORS\n"
        "CIN: U28129PN1979PLC141032\n"
        "DIN: 00135070\n"
        "SEBI Registration: INM000013004\n"
        "INR 7,100.00 million\n"
        "26,437,554 shares\n"
        "47.00%\n"
        "March 31, 2025\n"
        "PIN 411045\n"
        "Section 32\n"
        "Regulation 6(1)"
    )

    assert PIIDetector().detect(text) == []


def test_pin_alone_is_not_pii_but_address_with_pin_is() -> None:
    detector = PIIDetector()

    assert detector.detect("PIN 411045") == []
    entities = detector.detect("10 Example Road, Pune - 411045, Maharashtra, India")

    assert [entity.pii_type for entity in entities] == [PIIType.ADDRESS]


def test_diagnostics_are_count_only_and_consistent() -> None:
    text = "John Example Limited and John Example. Email: jane@example.com"

    result = PIIDetector().detect_with_diagnostics(text)

    assert result.diagnostics.raw_count == (
        result.diagnostics.accepted_count + result.diagnostics.rejected_count
    )
    assert result.diagnostics.counts_by_type_final["EMAIL"] == 1
    assert result.diagnostics.rejected_by_reason


def test_repeated_runs_are_deterministic() -> None:
    text = "Contact Person: Rahul Mehta Email: rahul@example.com"
    detector = PIIDetector()

    first = detector.detect(text)
    second = detector.detect(text)

    assert first == second


def test_cross_run_unified_detection_maps_to_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Email: john.")
    paragraph.add_run("doe@")
    paragraph.add_run("example.com")

    block = next(block for block in iter_text_blocks(document) if block.text)
    entities = PIIDetector().detect(block.text)
    email = next(entity for entity in entities if entity.pii_type == PIIType.EMAIL)
    run_slices = block.runs_for_span(email.start, email.end)

    assert email.text == "john.doe@example.com"
    assert [slice_.run_index for slice_ in run_slices] == [0, 1, 2]


def test_table_unified_detection() -> None:
    document = Document()
    table = document.add_table(rows=5, cols=1)
    table.cell(0, 0).text = "Contact Person: Ananya Reddy"
    table.cell(1, 0).text = "Company: Northstar Securities Limited"
    table.cell(2, 0).text = "Address: 201, Example Tower, Baner, Pune - 411045, Maharashtra, India"
    table.cell(3, 0).text = "Email: table@example.com"
    table.cell(4, 0).text = "Phone: +91 98765 43210"

    detector = PIIDetector()
    entities = [
        entity
        for block in iter_text_blocks(document)
        if block.source_type == SourceType.TABLE
        for entity in detector.detect(block.text)
    ]

    assert {
        PIIType.PERSON,
        PIIType.COMPANY,
        PIIType.ADDRESS,
        PIIType.EMAIL,
        PIIType.PHONE,
    }.issubset(set(types(entities)))


def test_header_footer_unified_detection_once() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Email: header@example.com"
    document.sections[0].footer.paragraphs[0].text = "Company: Summit Bank Limited"
    document.add_section(WD_SECTION.NEW_PAGE)

    block_results = PIIDetector().detect_blocks(iter_text_blocks(document))
    header_footer_entities = [
        entity
        for result in block_results
        for entity in result.entities
        if entity.pii_type in {PIIType.EMAIL, PIIType.COMPANY}
    ]

    assert Counter(entity.text for entity in header_footer_entities) == {
        "header@example.com": 1,
        "Summit Bank Limited": 1,
    }


def test_detect_many_uses_batch_path_and_returns_expected_shape() -> None:
    texts = ["Email: a@example.com", "Contact Person: Rahul Mehta"]

    results = PIIDetector().detect_many(texts)

    assert len(results) == 2
    assert [entity.pii_type for entity in results[0]] == [PIIType.EMAIL]
    assert [entity.pii_type for entity in results[1]] == [PIIType.PERSON]
