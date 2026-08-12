from collections.abc import Iterator
from pathlib import PurePosixPath
from typing import Any

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.document.models import DocumentStatistics, SourceType, TextBlock
from app.document.text_map import build_text_and_run_spans


def iter_text_blocks(document: DocxDocument) -> Iterator[TextBlock]:
    """Yield paragraph text blocks from body, tables, headers, and footers."""
    yield from _iter_body_blocks(document)
    yield from _iter_unique_header_footer_blocks(document, SourceType.HEADER)
    yield from _iter_unique_header_footer_blocks(document, SourceType.FOOTER)


def collect_document_statistics(document: DocxDocument) -> DocumentStatistics:
    blocks = list(iter_text_blocks(document))
    table_counts = _count_all_tables(document)

    return DocumentStatistics(
        section_count=len(document.sections),
        paragraph_count=len(blocks),
        table_count=table_counts["total"],
        text_block_count=len(blocks),
        body_block_count=sum(1 for block in blocks if block.source_type == SourceType.BODY),
        table_block_count=sum(
            1 for block in blocks if block.source_type == SourceType.TABLE
        ),
        header_block_count=sum(
            1 for block in blocks if block.source_type == SourceType.HEADER
        ),
        footer_block_count=sum(
            1 for block in blocks if block.source_type == SourceType.FOOTER
        ),
        header_part_count=len(_unique_header_footer_parts(document, SourceType.HEADER)),
        footer_part_count=len(_unique_header_footer_parts(document, SourceType.FOOTER)),
        nested_table_count=table_counts["nested"],
        embedded_media_count=_count_embedded_media(document),
    )


def _iter_body_blocks(document: DocxDocument) -> Iterator[TextBlock]:
    paragraph_index = 0
    table_index = 0
    part_id = "body:document"

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            yield _make_text_block(
                paragraph=paragraph,
                block_id=f"body:p:{paragraph_index:06d}",
                source_type=SourceType.BODY,
                part_id=part_id,
                location={"part": "body", "paragraph_index": paragraph_index},
            )
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            yield from _iter_table_blocks(
                table=table,
                source_prefix=f"body:tbl:{table_index:06d}",
                part_id=part_id,
                table_path=(table_index,),
                source_type=SourceType.TABLE,
                parent_location={"part": "body"},
            )
            table_index += 1


def _iter_unique_header_footer_blocks(
    document: DocxDocument, source_type: SourceType
) -> Iterator[TextBlock]:
    for part_id, container, location in _unique_header_footer_parts(
        document, source_type
    ):
        source_prefix = _part_block_prefix(source_type, part_id)
        for paragraph_index, paragraph in enumerate(container.paragraphs):
            yield _make_text_block(
                paragraph=paragraph,
                block_id=f"{source_prefix}:p:{paragraph_index:06d}",
                source_type=source_type,
                part_id=part_id,
                location={**location, "paragraph_index": paragraph_index},
            )

        for table_index, table in enumerate(container.tables):
            yield from _iter_table_blocks(
                table=table,
                source_prefix=f"{source_prefix}:tbl:{table_index:06d}",
                part_id=part_id,
                table_path=(table_index,),
                source_type=source_type,
                parent_location=location,
            )


def _unique_header_footer_parts(
    document: DocxDocument, source_type: SourceType
) -> list[tuple[str, Any, dict[str, Any]]]:
    seen_part_ids: set[str] = set()
    unique_parts: list[tuple[str, Any, dict[str, Any]]] = []
    variants = (
        ("default", "header" if source_type == SourceType.HEADER else "footer"),
        (
            "first_page",
            "first_page_header"
            if source_type == SourceType.HEADER
            else "first_page_footer",
        ),
        (
            "even_page",
            "even_page_header"
            if source_type == SourceType.HEADER
            else "even_page_footer",
        ),
    )

    for section_index, section in enumerate(document.sections):
        for variant, attr_name in variants:
            container = getattr(section, attr_name)
            # Several sections may point to the same Word part. Processing each
            # unique part once prevents future replacement passes from mutating
            # the same XML repeatedly.
            part_id = _part_id(source_type, container.part.partname)
            if part_id in seen_part_ids:
                continue
            seen_part_ids.add(part_id)
            unique_parts.append(
                (
                    part_id,
                    container,
                    {
                        "part": source_type.value.lower(),
                        "section_index": section_index,
                        "variant": variant,
                    },
                )
            )

    return unique_parts


def _iter_table_blocks(
    table: Table,
    source_prefix: str,
    part_id: str,
    table_path: tuple[int, ...],
    source_type: SourceType,
    parent_location: dict[str, Any],
) -> Iterator[TextBlock]:
    seen_cells: set[object] = set()

    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell_identity = cell._tc
            if cell_identity in seen_cells:
                continue
            seen_cells.add(cell_identity)

            cell_prefix = f"{source_prefix}:r:{row_index:03d}:c:{column_index:03d}"
            cell_location = {
                **parent_location,
                "table_path": table_path,
                "row_index": row_index,
                "column_index": column_index,
            }

            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield _make_text_block(
                    paragraph=paragraph,
                    block_id=f"{cell_prefix}:p:{paragraph_index:03d}",
                    source_type=source_type,
                    part_id=part_id,
                    location={**cell_location, "paragraph_index": paragraph_index},
                )

            for nested_index, nested_table in enumerate(cell.tables):
                yield from _iter_table_blocks(
                    table=nested_table,
                    source_prefix=f"{cell_prefix}:nt:{nested_index:03d}",
                    part_id=part_id,
                    table_path=(*table_path, row_index, column_index, nested_index),
                    source_type=source_type,
                    parent_location={
                        **cell_location,
                        "nested_table_index": nested_index,
                    },
                )


def _make_text_block(
    paragraph: Paragraph,
    block_id: str,
    source_type: SourceType,
    part_id: str,
    location: dict[str, Any],
) -> TextBlock:
    text, run_spans = build_text_and_run_spans(paragraph)
    return TextBlock(
        block_id=block_id,
        text=text,
        paragraph=paragraph,
        source_type=source_type,
        part_id=part_id,
        run_spans=run_spans,
        location=location,
    )


def _part_id(source_type: SourceType, partname: Any) -> str:
    return f"{source_type.value.lower()}:{partname}"


def _part_block_prefix(source_type: SourceType, part_id: str) -> str:
    part_name = PurePosixPath(part_id.split(":", 1)[1]).stem
    return f"{source_type.value.lower()}:{part_name}"


def _count_all_tables(document: DocxDocument) -> dict[str, int]:
    total = 0
    nested = 0

    def count_container(container: Any) -> None:
        nonlocal total, nested
        for table in container.tables:
            total += 1
            count_table(table, is_nested=False)

    def count_table(table: Table, is_nested: bool) -> None:
        nonlocal total, nested
        seen_cells: set[object] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_identity = cell._tc
                if cell_identity in seen_cells:
                    continue
                seen_cells.add(cell_identity)
                for nested_table in cell.tables:
                    total += 1
                    nested += 1
                    count_table(nested_table, is_nested=True)

    count_container(document)
    for source_type in (SourceType.HEADER, SourceType.FOOTER):
        for _, container, _ in _unique_header_footer_parts(document, source_type):
            count_container(container)

    return {"total": total, "nested": nested}


def _count_embedded_media(document: DocxDocument) -> int:
    package = document.part.package
    return sum(
        1
        for part in package.parts
        if str(part.partname).startswith("/word/media/")
    )
